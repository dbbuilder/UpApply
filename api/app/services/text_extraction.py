"""Text extraction service for attachments (PDF, DOCX, images via OCR)."""
import base64
import io
import logging
from typing import Tuple

from app.core.embeddings import get_openai_client

logger = logging.getLogger(__name__)


async def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        logger.error("pdfplumber not installed. Run: pip install pdfplumber")
        raise ImportError("pdfplumber is required for PDF extraction")

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    return "\n\n".join(text_parts)


async def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        raise ImportError("python-docx is required for DOCX extraction")

    doc = Document(io.BytesIO(file_bytes))
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n\n".join(text_parts)


async def extract_text_from_image(file_bytes: bytes, content_type: str = "image/png") -> str:
    """Extract text from an image using OpenAI Vision API (OCR)."""
    client = get_openai_client()

    # Encode image to base64
    base64_image = base64.b64encode(file_bytes).decode("utf-8")

    # Determine media type
    media_type = content_type if content_type.startswith("image/") else "image/png"

    try:
        response = await client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Extract ALL text visible in this image. Return only the extracted text content, preserving the structure and formatting as much as possible. If there is no text, return 'No text found in image.'",
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{media_type};base64,{base64_image}",
                                "detail": "high",
                            },
                        },
                    ],
                }
            ],
            max_tokens=4000,
        )

        return response.choices[0].message.content.strip()
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        raise


async def extract_text(
    file_bytes: bytes,
    content_type: str,
    filename: str = "",
) -> Tuple[str, str]:
    """Extract text from a file based on its content type.

    Args:
        file_bytes: The raw file bytes
        content_type: MIME type of the file
        filename: Original filename (used as fallback for type detection)

    Returns:
        Tuple of (extracted_text, extraction_status)
        extraction_status is one of: "success", "partial", "failed", "unsupported"
    """
    content_type_lower = content_type.lower()
    filename_lower = filename.lower()

    try:
        # PDF
        if "pdf" in content_type_lower or filename_lower.endswith(".pdf"):
            text = await extract_text_from_pdf(file_bytes)
            if text.strip():
                return text, "success"
            else:
                # PDF might be image-based, try OCR
                logger.info("PDF has no extractable text, attempting OCR...")
                text = await extract_text_from_image(file_bytes, "image/png")
                return text, "partial" if text else "failed"

        # DOCX
        if (
            "wordprocessingml" in content_type_lower
            or "msword" in content_type_lower
            or filename_lower.endswith(".docx")
            or filename_lower.endswith(".doc")
        ):
            text = await extract_text_from_docx(file_bytes)
            return text, "success" if text.strip() else "failed"

        # Images
        if content_type_lower.startswith("image/") or any(
            filename_lower.endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".gif", ".webp"]
        ):
            text = await extract_text_from_image(file_bytes, content_type)
            if text and "no text found" not in text.lower():
                return text, "success"
            else:
                return "", "partial"  # Image with no text is still valid

        # Plain text
        if "text/plain" in content_type_lower or filename_lower.endswith(".txt"):
            try:
                text = file_bytes.decode("utf-8")
                return text, "success"
            except UnicodeDecodeError:
                try:
                    text = file_bytes.decode("latin-1")
                    return text, "success"
                except Exception:
                    return "", "failed"

        # Unsupported type
        logger.warning(f"Unsupported file type: {content_type} ({filename})")
        return "", "unsupported"

    except ImportError as e:
        logger.error(f"Missing dependency for extraction: {e}")
        return "", "failed"
    except Exception as e:
        logger.error(f"Extraction failed for {filename}: {e}")
        return "", "failed"


async def extract_all_attachments(
    attachments: list[dict],
) -> Tuple[str, list[dict]]:
    """Extract text from multiple attachments.

    Args:
        attachments: List of dicts with keys: data (base64), filename, content_type

    Returns:
        Tuple of (combined_text, metadata_list)
        metadata_list contains extraction status for each file
    """
    all_text_parts = []
    metadata = []

    for attachment in attachments:
        filename = attachment.get("filename", "unknown")
        content_type = attachment.get("content_type", "application/octet-stream")
        data_b64 = attachment.get("data", "")

        try:
            # Decode base64
            file_bytes = base64.b64decode(data_b64)
            size = len(file_bytes)

            # Extract text
            text, status = await extract_text(file_bytes, content_type, filename)

            if text:
                all_text_parts.append(f"--- {filename} ---\n{text}")

            metadata.append({
                "filename": filename,
                "content_type": content_type,
                "size": size,
                "extraction_status": status,
                "error": None,
            })

        except Exception as e:
            logger.error(f"Failed to process attachment {filename}: {e}")
            metadata.append({
                "filename": filename,
                "content_type": content_type,
                "size": 0,
                "extraction_status": "failed",
                "error": str(e),
            })

    combined_text = "\n\n".join(all_text_parts)
    return combined_text, metadata
